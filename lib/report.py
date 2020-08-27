# -*- encoding: utf-8 -*-
'''
@File : hostscan.py
@Author : aydcyhr
@Contact : aydcyhr@gmail.com
'''

from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
import datetime
from  matplotlib import pyplot
import matplotlib.pyplot as plt
import socket
import os
import warnings

warnings.filterwarnings("ignore")

#数据
def reportfile(iplist,alivelist,portdic,vullist):
	try:
		document = Document()
		document.styles['Normal'].font.name = u'宋体'
		document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
		document.add_heading('hostscan 扫描报告', 0) #主标题
		document.add_heading('测试范围', level=1)   #一级标题
		#测试端口信息
		fw2= [str(i) for i in iplist]
		fw3 = ', '.join(fw2)
		document.add_paragraph(fw3)
		#存活数据信息
		document.add_heading('存活地址', level=1)   #一级标题
		ch2= [str(i) for i in alivelist]
		ch3 = '\n'.join(ch2)
		document.add_paragraph(ch3)
		#存活数据表格
		rantime=str(datetime.datetime.now().year) +str(datetime.datetime.now().month)+str(datetime.datetime.now().day)+str(datetime.datetime.now().hour)+str(datetime.datetime.now().minute)+str(datetime.datetime.now().second)
		png = './report/%s.png' %rantime
		document.add_heading('图表数据', level=1)   #一级标题
		document.add_heading("资产存活饼状图", level=3)
		#pyplot.rcParams['font.sans-serif']=['SimHei']
		pyplot.rcParams['axes.unicode_minus'] = False  # 用来正常显示负号
		pyplot.figure(figsize=(5, 4))
		alive = len(alivelist)
		unalive = len(iplist)-alive
		alive1 = "UP:"+str(alive)
		unalive1 = "Down:" +str(unalive)
		pyplot.title("ALive Data")
		pyplot.pie([alive,unalive], labels=[alive1, unalive1], startangle=90, counterclock=False)
		pyplot.savefig(png)
		document.add_picture(png)
		os.remove(png)
		#代言发
		rantime=str(datetime.datetime.now().year) +str(datetime.datetime.now().month)+str(datetime.datetime.now().day)+str(datetime.datetime.now().hour)+str(datetime.datetime.now().minute)+str(datetime.datetime.now().second)
		png = './report/%s.png' %rantime
		document.add_heading("Total Data", level=3)
		#plt.rcParams['font.sans-serif']=['SimHei']
		plt.rcParams['axes.unicode_minus'] = False  # 用来正常显示负号
		plt.figure(figsize=(5, 4))
		alive = len(alivelist)
		unalive = len(iplist)-alive
		ipnum = len(iplist)
		vulnum = len(vullist)
		plt.title("host data")
		
		y = [ipnum,alive,unalive,vulnum]
		x = ["Total", "UP" ,"Down", "Vuln"]
		plt.ylabel('Number')
		plt.bar(x, y)
		for x,y in enumerate(y):
			plt.text(x,y,'%s' %round(y,1),ha='center')
		plt.savefig(png)
		document.add_picture(png)
		os.remove(png)
		#端口开放详情
		document.add_heading("端口开放详情", level=1)
		newdic = []
		zidian = {}
		for ip in portdic:
			l = ip.split(":")[0]
			newdic.append(l)
			newdic = list(set(newdic))
		paixu = sorted(newdic,key=socket.inet_aton)

		for ip in paixu:
			canshu = []
			for port in portdic:
				l = port.split(":")[0]
				p = port.split(":")[-1]
				if l == ip:
					canshu.append(p)
			zidian[ip] = canshu   #zidian={'192.168.1.1': ['90', '901'], '192.168.1.11': ['10'], '192.168.1.13': ['930'], '192.168.2.11': ['910']}

		table = document.add_table(rows=1, cols=2,style='Normal Table')
		hdr_cells = table.rows[0].cells
		hdr_cells[0].text = 'IP'
		hdr_cells[1].text = '端口'
		for vl in zidian.keys():
			row_cells = table.add_row().cells
			row_cells[0].text = vl
			je2se = ''
			for i in zidian[vl]:
				je2se+= i+","
			row_cells[1].text =je2se 
		
 
	#漏洞部分
		document.add_heading("漏洞风险", level=1)
		if len(vullist)>0:
			for i in vullist:
				document.add_paragraph(i)
		else:
			document.add_paragraph("没有进行漏洞扫描或不存在漏洞")



	#声明部分
		document.add_heading("报告声明", level=1)
		p = document.add_paragraph('本报告由ADYHostScan工具扫描生成')
		p = document.add_paragraph('脚本地址：https://github.com/aydcyhr/ADYHostScan')


	#数据导出部分
		now=str(datetime.datetime.now().year) +str(datetime.datetime.now().month)+str(datetime.datetime.now().day)+str(datetime.datetime.now().hour)+str(datetime.datetime.now().minute)+str(datetime.datetime.now().second)

		path = './report/hostscan_%s.docx'  % now
		document.save(path)

		return path
	except :
		pass
  
